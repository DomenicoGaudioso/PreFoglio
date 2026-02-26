import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
#from kaleido import *
import io

# Importa le tue funzioni modificate
# Assicurati che def_PreFoglioPy.py sia nella stessa cartella
from def_PreFoglioPy import *
from def_ToPontiEC4 import *

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # <-- NUOVO: serve per centrare la didascalia

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import numpy as np
import pandas as pd

def crea_immagine_matplotlib(df_geom, df_res, piano, titolo, is_mobili=False):
    """Genera un grafico Matplotlib per Taglio o Momento. Se is_mobili=True disegna max (blu) e min (rosso)."""
    fig, ax = plt.subplots(figsize=(8, 4))
    
    # Gli assi geometrici dipendono ancora dal piano scelto per disegnare l'asta correttamente
    col_x, col_y = ('X', 'Z') if piano == 'XZ' else ('Y', 'Z')
    merged = pd.merge(df_geom, df_res, on='Elem', how='inner')

    # Sotto-funzione per non ripetere il calcolo geometrico del trapezio
    def disegna_poligono(x1, y1, x2, y2, v1, v2, colore):
        if pd.isna(v1): v1 = 0
        if pd.isna(v2): v2 = 0
        if abs(v1) >= 0.001 or abs(v2) >= 0.001:
            dx, dy = x2 - x1, y2 - y1
            L = np.sqrt(dx**2 + dy**2)
            if L > 0:
                nx, ny = -dy / L, dx / L
                ax.fill([x1, x1 + nx * v1, x2 + nx * v2, x2], 
                        [y1, y1 + ny * v1, y2 + ny * v2, y2], 
                        color=colore, alpha=0.4, edgecolor=colore)

    for _, row in merged.iterrows():
        x1, y1 = row[f'{col_x}_i'], row[f'{col_y}_i']
        x2, y2 = row[f'{col_x}_j'], row[f'{col_y}_j']
        
        # Disegna l'asta base
        ax.plot([x1, x2], [y1, y2], color='black', linewidth=1)

        if is_mobili:
            # Inviluppo: Disegna contemporaneamente il Max e il Min
            disegna_poligono(x1, y1, x2, y2, row.get('Val_I_max', 0), row.get('Val_J_max', 0), 'blue')
            disegna_poligono(x1, y1, x2, y2, row.get('Val_I_min', 0), row.get('Val_J_min', 0), 'red')
        else:
            # CDS standard: Disegna un solo trapezio
            v1, v2 = row.get('Val_I', 0), row.get('Val_J', 0)
            colore = 'red' if (v1 + v2) / 2 < 0 else 'blue'
            disegna_poligono(x1, y1, x2, y2, v1, v2, colore)

    ax.set_title(titolo)
    ax.invert_yaxis()
    ax.grid(True, linestyle='--', alpha=0.5)
    
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

def genera_report_word_matplotlib(dfs, geom, piano):
    """Genera il documento Word inserendo SOLO le immagini e le relative didascalie."""
    doc = Document()
    
    # FISSO: Calcoliamo e mostriamo solo Taglio Vz e Momento My
    grandezze = ['Shear-z', 'Moment-y']

    for sheet_name in ['CDS', 'Mobili']:
        if sheet_name not in dfs: continue
            
        df_data = dfs[sheet_name].copy() 
        if 'Load' not in df_data.columns: continue

        # --- GESTIONE MOBILI ---
        if sheet_name == 'Mobili':
            # CORRETTO: (?i) all'inizio dell'espressione regolare
            df_data['Base_Load'] = df_data['Load'].astype(str).str.replace(r'(?i)\s*\((max|min)\)', '', regex=True)
            
            for base_load in df_data['Base_Load'].unique():
                df_base = df_data[df_data['Base_Load'] == base_load]
                if df_base.empty: continue
                
                for val_col in grandezze:
                    if val_col not in df_base.columns: continue
                    
                    if 'Component' in df_base.columns:
                        df_comp = df_base[df_base['Component'] == val_col]
                    else:
                        df_comp = df_base
                        
                    if df_comp.empty: continue
                    
                    # CORRETTO: (?i) all'inizio dell'espressione regolare
                    df_max = df_comp[df_comp['Load'].astype(str).str.contains(r'(?i)\(max\)', regex=True)]
                    df_min = df_comp[df_comp['Load'].astype(str).str.contains(r'(?i)\(min\)', regex=True)]
                    
                    res_list = []
                    elems = set(df_comp['Elem'].unique())
                    
                    for elem in elems:
                        elem_max = df_max[df_max['Elem'] == elem]
                        elem_min = df_min[df_min['Elem'] == elem]
                        
                        # Max
                        if not elem_max.empty:
                            i_max = elem_max[elem_max['Part'].astype(str).str.startswith('I', na=False)][val_col]
                            j_max = elem_max[elem_max['Part'].astype(str).str.startswith('J', na=False)][val_col]
                            if i_max.empty: i_max = elem_max[val_col]
                            if j_max.empty: j_max = elem_max[val_col]
                            v_i_max = i_max.iloc[0] if not i_max.empty else 0
                            v_j_max = j_max.iloc[-1] if not j_max.empty else 0
                        else:
                            v_i_max = v_j_max = 0
                            
                        # Min
                        if not elem_min.empty:
                            i_min = elem_min[elem_min['Part'].astype(str).str.startswith('I', na=False)][val_col]
                            j_min = elem_min[elem_min['Part'].astype(str).str.startswith('J', na=False)][val_col]
                            if i_min.empty: i_min = elem_min[val_col]
                            if j_min.empty: j_min = elem_min[val_col]
                            v_i_min = i_min.iloc[0] if not i_min.empty else 0
                            v_j_min = j_min.iloc[-1] if not j_min.empty else 0
                        else:
                            v_i_min = v_j_min = 0
                            
                        res_list.append({
                            'Elem': elem,
                            'Val_I_max': v_i_max, 'Val_J_max': v_j_max,
                            'Val_I_min': v_i_min, 'Val_J_min': v_j_min
                        })
                        
                    if not res_list: continue
                    res_extract = pd.DataFrame(res_list)
                    
                    # Genera l'immagine (titolo interno al grafico)
                    titolo_plot = f"{val_col} Inviluppo"
                    img_buffer = crea_immagine_matplotlib(geom, res_extract, piano, titolo_plot, is_mobili=True)
                    
                    # --- INSERIMENTO NEL WORD ---
                    doc.add_picture(img_buffer, width=Inches(6.0))
                    
                    # Inserisce la didascalia sotto e la centra
                    testo_didascalia = f"Figura: {sheet_name} - {base_load} - {val_col} (Inviluppo)"
                    p = doc.add_paragraph(testo_didascalia)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("")

        # --- GESTIONE CDS ---
        else:
            for load in df_data['Load'].unique():
                df_load = df_data[df_data['Load'] == load]
                if df_load.empty: continue
                
                for val_col in grandezze:
                    if val_col not in df_load.columns: continue
                    
                    grouped = df_load.groupby('Elem')[val_col]
                    res_extract = pd.DataFrame({
                        'Val_I': grouped.first(),
                        'Val_J': grouped.last()
                    }).reset_index()
                    
                    # Genera l'immagine
                    titolo_plot = f"{val_col}"
                    img_buffer = crea_immagine_matplotlib(geom, res_extract, piano, titolo_plot, is_mobili=False)
                    
                    # --- INSERIMENTO NEL WORD ---
                    doc.add_picture(img_buffer, width=Inches(6.0))
                    
                    # Inserisce la didascalia sotto e la centra
                    testo_didascalia = f"Figura: {sheet_name} - {load} - {val_col}"
                    p = doc.add_paragraph(testo_didascalia)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("")

    word_io = io.BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io


st.set_page_config(page_title="GaudiCose - FEM Viewer", layout="wide")

# --- FUNZIONI DI SUPPORTO ---

@st.cache_data
def load_excel_sheets(uploaded_file):
    """Carica l'Excel una sola volta e lo tiene in memoria come dizionario."""
    return pd.read_excel(uploaded_file, sheet_name=None)

def to_excel_bytes(df):
    """Converte DF in bytes per il download."""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False)
    return output.getvalue()

def calcola_diagrammi(df_geom, df_res, piano, componente):
    """
    Funzione Core per generare i trapezi di carico.
    df_geom: Merge di Element e Nodi (contiene Xi, Yi, Xj, Yj)
    df_res: Risultati (Valore a I e Valore a J per ogni elem)
    """
    traces = []
    
    # Mappatura assi in base al piano scelto
    if piano == 'XZ':
        col_x, col_y = 'X', 'Z'
    else: # YZ
        col_x, col_y = 'Y', 'Z'

    # Iteriamo per ogni elemento geometrico che ha risultati
    merged = pd.merge(df_geom, df_res, on='Elem', how='inner')

    for _, row in merged.iterrows():
        # Coordinate Nodali
        x1, y1 = row[f'{col_x}_i'], row[f'{col_y}_i']
        x2, y2 = row[f'{col_x}_j'], row[f'{col_y}_j']
        
        # Valori Sollecitazione
        v1 = row['Val_I']
        v2 = row['Val_J']
        
        # Se nullo sostituisci con 0
        if pd.isna(v1): v1 = 0
        if pd.isna(v2): v2 = 0

        # 1. Disegna l'asta (Linea Strutturale)
        traces.append(go.Scatter(
            x=[x1, x2], y=[y1, y2],
            mode='lines',
            line=dict(color='black', width=1),
            showlegend=False,
            hoverinfo='skip'
        ))

        # Se i valori sono prossimi a zero, non disegniamo il trapezio
        if abs(v1) < 0.001 and abs(v2) < 0.001:
            continue

        # 2. Calcolo Geometria Trapezio
        dx = x2 - x1
        dy = y2 - y1
        L = np.sqrt(dx**2 + dy**2)
        
        if L > 0:
            # Vettore Normale unitario (rotazione 90 gradi)
            nx = -dy / L
            ny = dx / L
            
            # Coordinate punti diagramma
            # P1(I) -> P2(I_scalato) -> P3(J_scalato) -> P4(J) -> Chiudi
            
            p1_x, p1_y = x1, y1
            p2_x, p2_y = x1 + nx * v1 , y1 + ny * v1 
            p3_x, p3_y = x2 + nx * v2 , y2 + ny * v2
            p4_x, p4_y = x2, y2
            
            # Colore (Rosso neg, Blu pos basato sulla media)
            val_medio = (v1 + v2) / 2
            fill_color = 'rgba(255, 0, 0, 0.4)' if val_medio < 0 else 'rgba(0, 0, 255, 0.4)'
            border_color = 'red' if val_medio < 0 else 'blue'

            traces.append(go.Scatter(
                x=[p1_x, p2_x, p3_x, p4_x],
                y=[p1_y, p2_y, p3_y, p4_y],
                fill='toself',
                fillcolor=fill_color,
                mode='lines',
                line=dict(color=border_color, width=1),
                name=f'Elem {row["Elem"]}',
                text=f"Elem: {row['Elem']}<br>I: {v1:.2f}<br>J: {v2:.2f}",
                hoverinfo='text'
            ))
            
    return traces

# --- INTERFACCIA ---
st.title("🏗️ CDSEditing")

uploaded_file = st.sidebar.file_uploader("Carica Excel Input", type=["xlsx"])

if uploaded_file:
    # Caricamento Excel in Dizionario
    with st.spinner("Lettura file..."):
        dfs = load_excel_sheets(uploaded_file)
    st.sidebar.success("File Caricato!")

    # Check Fogli necessari
    required_sheets = ['Point', 'Element']
    missing = [s for s in required_sheets if s not in dfs]
    
    if missing:
        st.error(f"Mancano i fogli fondamentali nell'Excel: {missing}")
    else:
        # Prepara Geometria Base (Nodi + Elementi)
        df_elem = dfs['Element'][['Element', 'Node1', 'Node2']].rename(columns={'Element': 'Elem'})
        df_node = dfs['Point'][['Node', 'X', 'Y', 'Z']]
        
        # Merge Coordinate I
        geom = pd.merge(df_elem, df_node, left_on='Node1', right_on='Node').rename(
            columns={'X':'X_i', 'Y':'Y_i', 'Z':'Z_i'}).drop(columns=['Node'])
        # Merge Coordinate J
        geom = pd.merge(geom, df_node, left_on='Node2', right_on='Node').rename(
            columns={'X':'X_j', 'Y':'Y_j', 'Z':'Z_j'}).drop(columns=['Node'])

        # --- TABS ---
        tab_exp, tab_view, tab_group = st.tabs(["💾 Generazione File", "📐 Visualizzazione Grafica", " work data"])

# --- TAB 1: VISUALIZZAZIONE ---
        with tab_view:
            row1_col1, row1_col2, row1_col3 = st.columns(3)
            row2_col1, row2_col2, row2_col3 = st.columns(3)

            # 1. Scelta Foglio Dati (CDS o Mobili)
            data_sheets = [k for k in dfs.keys() if k not in ['Point', 'Element', 'Section']]
            sel_sheet = row1_col1.selectbox("Sorgente Dati", data_sheets, index=0)
            
            df_data = dfs[sel_sheet]
            
            # --- LOGICA DI FILTRO DINAMICA ---
            # Caso A: Foglio "Mobili" (ha colonna 'Component')
            if 'Component' in df_data.columns:
                # 2. Scelta Load
                loads = df_data['Load'].unique()
                sel_load = row1_col2.selectbox("Caso di Carico (Load)", loads)
                
                # 3. Scelta Componente Dominante (es. voglio la condizione di Max Moment-y)
                comps_avail = df_data['Component'].unique()
                # Cerchiamo di preselezionare Moment-y se c'è
                idx_comp = list(comps_avail).index('Moment-y') if 'Moment-y' in comps_avail else 0
                sel_criterion = row2_col1.selectbox("Criterio Inviluppo (Component)", comps_avail, index=idx_comp)
                
                # 4. Scelta Grandezza da Plottare
                # Di solito plottiamo la stessa grandezza del criterio, ma potremmo voler vedere le concomitanti
                cols_num = df_data.select_dtypes(include=np.number).columns.tolist()
                cols_exclude = ['Elem', 'Node', 'Material', 'Section', 'Part']
                cols_plot = [c for c in cols_num if c not in cols_exclude]
                
                # Preselezioniamo la stessa del criterio se esiste, altrimenti la prima
                idx_plot = cols_plot.index(sel_criterion) if sel_criterion in cols_plot else 0
                sel_val_col = row2_col2.selectbox("Grandezza da Plottare", cols_plot, index=idx_plot)

                # FILTRO EFFETTIVO PER MOBILI
                df_filtered = df_data[
                    (df_data['Load'] == sel_load) & 
                    (df_data['Component'] == sel_criterion)
                ]
                st.caption(f"Visualizzando: **{sel_val_col}** quando **{sel_criterion}** è massimizzato/minimizzato.")

            # Caso B: Foglio "CDS" (senza colonna 'Component')
            else:
                # 2. Scelta Load
                if 'Load' in df_data.columns:
                    loads = df_data['Load'].unique()
                    sel_load = row1_col2.selectbox("Caso di Carico", loads)
                    df_filtered = df_data[df_data['Load'] == sel_load]
                else:
                    df_filtered = df_data # Nessun filtro load
                
                # 3. Scelta Grandezza
                cols_num = df_filtered.select_dtypes(include=np.number).columns.tolist()
                cols_exclude = ['Elem', 'Node', 'Material', 'Section', 'Part', 'Step']
                cols_plot = [c for c in cols_num if c not in cols_exclude]
                
                idx_def = cols_plot.index('Moment-y') if 'Moment-y' in cols_plot else 0
                sel_val_col = row2_col1.selectbox("Grandezza da Plottare", cols_plot, index=idx_def)
                
                # Placeholder per mantenere l'allineamento
                row2_col2.write("") 

            # Parametri Grafici
            sel_plane = row1_col3.selectbox("Piano di Vista", ["XZ", "YZ"])

            # --- GENERAZIONE PLOT ---
            if df_filtered.empty:
                st.warning("Nessun dato trovato con i filtri selezionati.")
            else:
                with st.spinner("Elaborazione grafico..."):
                    # Estrazione valori I e J per ogni elemento
                    # Assunzione: Il file è ordinato per Elem, e le righe Part rappresentano le sezioni lungo l'asta.
                    # Prendiamo il Primo valore (I) e l'Ultimo (J) del gruppo elemento.
                    
                    grouped = df_filtered.groupby('Elem')[sel_val_col]
                    res_extract = pd.DataFrame({
                        'Val_I': grouped.first(),
                        'Val_J': grouped.last()
                    }).reset_index()

                    # Calcolo tracce
                    traces = calcola_diagrammi(geom, res_extract, sel_plane, sel_val_col)

                    # Plot
                    fig = go.Figure(data=traces)
                    fig.update_layout(
                        title=f"Diagramma {sel_val_col} su {sel_plane}",
                        height=800,
                        showlegend=False,
                        plot_bgcolor="white",
                        hovermode="closest"
                    )
                    fig.update_yaxes(autorange="reversed")
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Inverti asse Y (negativo sopra, positivo sotto)
                         
# --- SEZIONE EXPORT WORD CON MATPLOTLIB ---
            st.divider()
            st.subheader("📄 Generazione Report CDS")
            st.write("Genera un documento Word con i grafici di Taglio e Momento.")
            
            if st.button("Genera Report Word"):
                with st.spinner("Creazione dei grafici in corso... (L'operazione richiede qualche secondo)"):
                    try:
                        word_file = genera_report_word_matplotlib(dfs, geom, sel_plane)
                        
                        st.download_button(
                            label="📥 Scarica Report_Sollecitazioni.docx",
                            data=word_file,
                            file_name=f"Report_Sollecitazioni_{sel_plane}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success("Report Word generato! Clicca il pulsante per scaricarlo.")
                    except Exception as e:
                        st.error(f"Errore durante l'esportazione Word: {e}")
                    


        # TAB 2: EXPORT
        with tab_exp:
            st.write("Usa i pulsanti per ottenere le massime e minimi sollecitazioni per ogni concio")
            
            st.subheader("Export 1 (SLU)", divider=True)
            if st.button("Export Standard"):
                try:
                    # Passiamo 'dfs' come input_data
                    res1 = Run_Export1Out_SuperFoglio(dfs)
                    #st.write(res1)
                    st.download_button(
                        "📥 Scarica Excel 1",
                        data=res1,
                        file_name="Output_1.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                    st.success("Fatto.")
                except Exception as e:
                    st.error(f"Errore: {e}")

            st.subheader("Export 2 (Fatica)", divider=True)
            metodo = st.selectbox("Metodo", [1, 2])
            if st.button("Export Fatica"):
                try:
                    res2 = Run_Export2Out_SuperFoglio(dfs, metodo=metodo)
                    
                    st.download_button(
                        "📥 Scarica Excel 2",
                        data=res2,
                        file_name="Output_2.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                    st.success("Fatto.")
                except Exception as e:
                    st.error(f"Errore: {e}")
                    
        # --- TAB 3: LAVORO CON I DATI PER RIESPORTARLI ---
        with tab_group:
            print("ciao")
            # 1. Lettura dei dati
            elements_df = dfs["Element"]

            # 2. Ordinamento per Proprietà (sezione) e ID Elemento
            elements_df = elements_df.sort_values(by=['Property', 'Element']).reset_index(drop=True)

            # 3. Calcolo della differenza tra l'ID corrente e il precedente all'interno della stessa Proprietà
            elements_df['Diff'] = elements_df.groupby('Property')['Element'].diff()

            # 4. Un nuovo gruppo inizia quando la differenza di ID non è 1 
            # (il primo elemento di ogni Proprietà avrà Diff = NaN, che conta come True nel test != 1)
            elements_df['NewGroup'] = (elements_df['Diff'] != 1).astype(int)

            # 5. La somma cumulativa identifica un numero di sottogruppo incrementale per i blocchi contigui
            elements_df['SubGroup'] = elements_df.groupby('Property')['NewGroup'].cumsum()

            # 6. Creazione del nome finale del gruppo
            elements_df['GroupName'] = 'group' + elements_df['Property'].astype(str) + '_' + elements_df['SubGroup'].astype(str)

            # 7. Salvataggio del risultato
            result_df = elements_df[['Element', 'Property', 'Node1', 'Node2', 'GroupName']]
            #result_df.to_csv("Element_Groups.csv", index=False)
            st.write(result_df)
            
            # Load the previously generated groups and the points
            # Ensure groups are sorted by Element to maintain the sequence
            groups_df = result_df.sort_values(by=['GroupName', 'Element'])
            points_df = df_node

            # Function to get I, K, J nodes for a group
            def get_key_nodes(group):
                # The nodes in order. 
                # Since elements are contiguous, Node2 of elem i is Node1 of elem i+1.
                # The full sequence of nodes is Node1 of all elements + Node2 of the last element.
                nodes = group['Node1'].tolist() + [group['Node2'].iloc[-1]]
                
                n_nodes = len(nodes)
                node_I = nodes[0]
                node_J = nodes[-1]
                node_K = nodes[n_nodes // 2] # Middle node
                
                return pd.Series({'Node_I': node_I, 'Node_K': node_K, 'Node_J': node_J})
            
            # Apply to each group
            group_nodes = groups_df.groupby('GroupName').apply(get_key_nodes).reset_index()

            # Now we need to get the coordinates for these nodes
            # Merge for Node I
            merged_I = pd.merge(group_nodes, points_df, left_on='Node_I', right_on='Node', how='left')
            merged_I = merged_I.rename(columns={'X': 'X_I', 'Y': 'Y_I', 'Z': 'Z_I'}).drop(columns=['Node'])

            # Merge for Node K
            merged_K = pd.merge(merged_I, points_df, left_on='Node_K', right_on='Node', how='left')
            merged_K = merged_K.rename(columns={'X': 'X_K', 'Y': 'Y_K', 'Z': 'Z_K'}).drop(columns=['Node'])

            # Merge for Node J
            final_df = pd.merge(merged_K, points_df, left_on='Node_J', right_on='Node', how='left')
            final_df = final_df.rename(columns={'X': 'X_J', 'Y': 'Y_J', 'Z': 'Z_J'}).drop(columns=['Node'])

            # Sort properly by Property and SubGroup
            final_df['Property_Num'] = final_df['GroupName'].str.extract(r'group(\d+)_').astype(int)
            final_df['SubGroup_Num'] = final_df['GroupName'].str.extract(r'_(\d+)').astype(int)
            final_df = final_df.sort_values(['Property_Num', 'SubGroup_Num']).drop(columns=['Property_Num', 'SubGroup_Num']).reset_index(drop=True)

            # Save to CSV
            output_file = "Group_Coordinates"
            st.write(final_df)
            
            # SOLLECITAZIONI
            cds_df = dfs["CDS"]
            mobili_df = dfs["Mobili"]

            #print("CDS columns:", cds_df.columns.tolist())
            #print("Mobili columns:", mobili_df.columns.tolist())

            #print("\nCDS head:")
            #print(cds_df.head(2))

            #print("\nMobili head:")
            #print(mobili_df.head(2))
            
            # Create a mapping from GroupName to its elements, first element, and last element
            group_info = groups_df.groupby('GroupName').agg(
                Elements=('Element', list),
                First_Element=('Element', 'first'),
                Last_Element=('Element', 'last')
            ).reset_index()
            
            # Extract Property and Subgroup for sorting
            group_info['Property_Num'] = group_info['GroupName'].str.extract(r'group(\d+)_').astype(int)
            group_info['SubGroup_Num'] = group_info['GroupName'].str.extract(r'_(\d+)').astype(int)
            
            
            # Ensure forces are numeric
            force_cols = ['Axial', 'Shear-y', 'Shear-z', 'Torsion', 'Moment-y', 'Moment-z']
            for col in force_cols:
                cds_df[col] = pd.to_numeric(cds_df[col], errors='coerce').fillna(0)
                mobili_df[col] = pd.to_numeric(mobili_df[col], errors='coerce').fillna(0)

            # Function to get extreme value (max absolute value, keeping sign)
            def get_extreme(series):
                # Get index of max absolute value
                if len(series) == 0:
                    return 0
                idx = series.abs().idxmax()
                return series.loc[idx]

            def process_forces(df, group_info, has_component=False):
                # Join with groups_df to get GroupName for each element
                df_merged = df.merge(groups_df[['Element', 'GroupName']], left_on='Elem', right_on='Element')
                
                groupby_cols = ['GroupName', 'Load', 'Component'] if has_component else ['GroupName', 'Load']
                
                results = []
                
                for name, group_data in df_merged.groupby(groupby_cols):
                    g_name = name[0]
                    # Get first and last element of this group
                    g_info = group_info[group_info['GroupName'] == g_name].iloc[0]
                    first_el = g_info['First_Element']
                    last_el = g_info['Last_Element']
                    
                    # I forces: Elem == first_el and Part starts with 'I['
                    i_data = group_data[(group_data['Elem'] == first_el) & (group_data['Part'].str.startswith('I['))]
                    
                    # J forces: Elem == last_el and Part starts with 'J['
                    j_data = group_data[(group_data['Elem'] == last_el) & (group_data['Part'].str.startswith('J['))]
                    
                    res_row = list(name)
                    
                    # Append I forces
                    for col in force_cols:
                        val = i_data[col].values[0] if len(i_data) > 0 else np.nan
                        res_row.append(val)
                        
                    # Append K (extreme) forces
                    for col in force_cols:
                        val = get_extreme(group_data[col])
                        res_row.append(val)
                        
                    # Append J forces
                    for col in force_cols:
                        val = j_data[col].values[0] if len(j_data) > 0 else np.nan
                        res_row.append(val)
                        
                    results.append(res_row)
                    
                cols = groupby_cols + [f"{c}_I" for c in force_cols] + [f"{c}_K" for c in force_cols] + [f"{c}_J" for c in force_cols]
                res_df = pd.DataFrame(results, columns=cols)
                    
                # Sort logically
                res_df['Property_Num'] = res_df['GroupName'].str.extract(r'group(\d+)_').astype(int)
                res_df['SubGroup_Num'] = res_df['GroupName'].str.extract(r'_(\d+)').astype(int)
                res_df = res_df.sort_values(['Property_Num', 'SubGroup_Num', 'Load']).drop(columns=['Property_Num', 'SubGroup_Num']).reset_index(drop=True)
                
                return res_df

            cds_res = process_forces(cds_df, group_info, has_component=False)
            mobili_res = process_forces(mobili_df, group_info, has_component=True)

            # Save to CSV
            #cds_res.to_csv("Sollecitazioni_CDS_Processed.csv", index=False)
            #mobili_res.to_csv("Sollecitazioni_Mobili_Processed.csv", index=False)

            print("CDS head:")
            st.write(cds_res)
            print("Mobili head:")
            st.write(mobili_res)
            
            ## CREAZIONE FILE PONTI EC4
            
            button_ec4 = st.button("esporta EC4")
            button_mix = st.button("esporta mixBridge")
            
            if button_ec4:
                #final_df #modello ridotto
                
                df_elem = dfs['Element'][['Element', 'Node1', 'Node2']].rename(columns={'Element': 'Elem'})
                df_node = dfs['Point'][['Node', 'X', 'Y', 'Z']]
                df_secEC4 = dfs['sectionEC4'] 
        
                try:
                    # 1. Genera il dizionario delle sezioni
                    Model_conci = ModelConci_AddSection(final_df, df_secEC4)
                    
                    # 2. Genera il testo del file .bak (senza scriverlo su disco)
                    file_content = wPontiEC4_Model(final_df, Model_conci)
                    
                    # 3. Crea il pulsante di download per il modello
                    st.download_button(
                        label="⬇️ Scarica il Modello EC4",
                        data=file_content,
                        file_name="modelEC4.txt",
                        mime="text/plain"
                    )
                    st.success("Modello EC4 generato! Clicca il pulsante qui sopra per scaricarlo.")
                except Exception as e:
                    st.error(f"Errore generazione Modello: {e}")
                    
                try:
                    # Aggiunti PathOut e nameFile per permettere il salvataggio
                    comb_PontiEC4(final_df, pathCDS, PathOut, nameFile)
                    st.success("Sollecitazioni EC4 generate con successo!")
                except Exception as e:
                    st.error(f"Errore generazione Sollecitazioni: {e}")
                          

else:
    st.info("Attesa caricamento file Excel...")